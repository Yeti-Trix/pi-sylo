"""Read-only DOCX walk: paragraphs, headings, tables, and image anchors.

The document body is walked in order (paragraphs and tables interleaved),
so every image reference keeps its position in the text plus nearby caption
candidates. Media names (``image3.png``) are the shared key between
``read_docx`` markers and ``extract_docx_images`` output files.
"""

from __future__ import annotations

import re
import tempfile
import zipfile
from contextlib import contextmanager
from pathlib import Path
from typing import Any, Iterator

from fix_strict_ooxml import convert_docx, is_strict_ooxml

DOCX_MEDIA_PREFIX = "word/media/"
IMAGE_EXTENSIONS = frozenset(
    {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".emf", ".wmf", ".webp"}
)
NON_VISION_FORMATS = frozenset({".emf", ".wmf", ".tif", ".tiff"})

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS_V = "urn:schemas-microsoft-com:vml"

CAPTION_RE = re.compile(r"^\s*(figure|fig\.?|table|photo|image|diagram)\s*\.?\s*\d", re.IGNORECASE)
HEADING_STYLE_RE = re.compile(r"^heading\s+(\d)$", re.IGNORECASE)


@contextmanager
def open_readable(docx_path: Path) -> Iterator[tuple[Any, bool]]:
    """Yield ``(Document, converted_from_strict)``; never mutates the source file.

    Strict OOXML files are converted to a temp copy first (python-docx cannot
    open Strict packages).
    """
    from docx import Document

    docx_path = docx_path.resolve()
    if not docx_path.is_file():
        raise ValueError(f"DOCX not found: {docx_path}")
    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"Not a .docx file: {docx_path}")

    if is_strict_ooxml(docx_path):
        with tempfile.TemporaryDirectory() as tmp:
            fixed = convert_docx(docx_path, Path(tmp) / "transitional.docx")
            yield Document(str(fixed)), True
    else:
        yield Document(str(docx_path)), False


def _style_name(para: Any) -> str:
    try:
        return str(para.style.name or "")
    except Exception:  # noqa: BLE001 — corrupt style refs should not kill the read
        return ""


def _heading_level(style_name: str) -> int | None:
    if style_name.lower() == "title":
        return 0
    m = HEADING_STYLE_RE.match(style_name)
    return int(m.group(1)) if m else None


def _media_name_for_rid(doc: Any, rid: str) -> str | None:
    rel = doc.part.rels.get(rid)
    if rel is None or getattr(rel, "is_external", False):
        return None
    target = str(rel.target_ref)
    if "media/" not in target.replace("\\", "/"):
        return None
    return Path(target).name


def _paragraph_images(doc: Any, para: Any) -> list[dict[str, Any]]:
    """Image refs in one paragraph: DrawingML pictures and legacy VML shapes."""
    found: list[dict[str, Any]] = []
    el = para._p  # noqa: SLF001 — python-docx has no public drawing API

    for blip in el.findall(f".//{{{NS_A}}}blip"):
        rid = blip.get(f"{{{NS_R}}}embed") or blip.get(f"{{{NS_R}}}link")
        if not rid:
            continue
        media = _media_name_for_rid(doc, rid)
        if media:
            found.append({"media_name": media})

    for imagedata in el.findall(f".//{{{NS_V}}}imagedata"):
        rid = imagedata.get(f"{{{NS_R}}}id")
        if not rid:
            continue
        media = _media_name_for_rid(doc, rid)
        if media and not any(f["media_name"] == media for f in found):
            found.append({"media_name": media})

    if found:
        alts = [
            (dp.get("descr") or dp.get("name") or "").strip()
            for dp in el.findall(f".//{{{NS_WP}}}docPr")
        ]
        alts = [a for a in alts if a]
        for i, ref in enumerate(found):
            if i < len(alts):
                ref["alt_text"] = alts[i]
    return found


def walk_blocks(doc: Any) -> list[dict[str, Any]]:
    """Body blocks in document order: paragraph and table dicts.

    Each block: ``{"index", "kind": "paragraph"|"table", ...}``. Paragraphs
    carry ``text``, ``style``, ``images``; tables carry ``rows``, ``cols``,
    ``header``, ``images`` (from cell paragraphs).
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    blocks: list[dict[str, Any]] = []
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == f"{{{NS_W}}}p":
            para = Paragraph(child, doc)
            blocks.append(
                {
                    "index": len(blocks),
                    "kind": "paragraph",
                    "text": para.text or "",
                    "style": _style_name(para),
                    "images": _paragraph_images(doc, para),
                }
            )
        elif tag == f"{{{NS_W}}}tbl":
            table = Table(child, doc)
            rows = len(table.rows)
            cols = len(table.columns) if rows else 0
            header: list[str] = []
            images: list[dict[str, Any]] = []
            char_count = 0
            for r_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    for cp in cell.paragraphs:
                        char_count += len(cp.text or "")
                        images.extend(_paragraph_images(doc, cp))
                    if r_idx == 0:
                        header.append(cell.text.strip())
            blocks.append(
                {
                    "index": len(blocks),
                    "kind": "table",
                    "rows": rows,
                    "cols": cols,
                    "header": header,
                    "char_count": char_count,
                    "images": images,
                }
            )
    return blocks


def guess_title(doc: Any, blocks: list[dict[str, Any]], docx_path: Path) -> str:
    try:
        core_title = (doc.core_properties.title or "").strip()
        if core_title:
            return core_title
    except Exception:  # noqa: BLE001
        pass
    for block in blocks:
        if block["kind"] != "paragraph":
            continue
        level = _heading_level(block["style"])
        text = block["text"].strip()
        if level is not None and text:
            return text
    for block in blocks:
        if block["kind"] == "paragraph" and block["text"].strip():
            return block["text"].strip()[:120]
    return docx_path.stem


def collect_headings(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    headings: list[dict[str, Any]] = []
    for block in blocks:
        if block["kind"] != "paragraph":
            continue
        level = _heading_level(block["style"])
        text = block["text"].strip()
        if level is not None and level >= 1 and text:
            headings.append({"level": level, "title": text, "index": block["index"]})
    return headings


def _marker(ref: dict[str, Any]) -> str:
    alt = ref.get("alt_text")
    return f"[image: {ref['media_name']}" + (f" - {alt}]" if alt else "]")


def read_docx_summary(
    docx_path: Path,
    *,
    max_paragraphs: int = 400,
    offset: int = 0,
) -> dict[str, Any]:
    """Structured read: title guess, ordered paragraphs with image markers, headings, tables."""
    with open_readable(docx_path) as (doc, converted):
        blocks = walk_blocks(doc)

        char_count = 0
        image_refs: list[dict[str, Any]] = []
        for block in blocks:
            if block["kind"] == "paragraph":
                char_count += len(block["text"])
            else:
                char_count += block["char_count"]
            for ref in block["images"]:
                entry = {"media_name": ref["media_name"], "block_index": block["index"]}
                if ref.get("alt_text"):
                    entry["alt_text"] = ref["alt_text"]
                image_refs.append(entry)

        paragraphs: list[dict[str, Any]] = []
        tables_summary: list[dict[str, Any]] = []
        truncated = False
        next_offset: int | None = None
        for block in blocks:
            if block["index"] < offset:
                continue
            if block["kind"] == "table":
                tables_summary.append(
                    {
                        "index": block["index"],
                        "rows": block["rows"],
                        "cols": block["cols"],
                        "header": block["header"],
                        **(
                            {"images": [r["media_name"] for r in block["images"]]}
                            if block["images"]
                            else {}
                        ),
                    }
                )
                continue
            text = block["text"].strip()
            if not text and not block["images"]:
                continue
            if len(paragraphs) >= max_paragraphs:
                truncated = True
                next_offset = block["index"]
                break
            display = text
            if block["images"]:
                markers = " ".join(_marker(r) for r in block["images"])
                display = f"{text} {markers}".strip()
            entry: dict[str, Any] = {"index": block["index"], "text": display}
            style = block["style"]
            if style and style.lower() not in {"normal", "default paragraph font"}:
                entry["style"] = style
            paragraphs.append(entry)

        result: dict[str, Any] = {
            "docx_path": str(docx_path.resolve()),
            "converted_from_strict": converted,
            "title_guess": guess_title(doc, blocks, docx_path),
            "char_count": char_count,
            "block_count": len(blocks),
            "table_count": sum(1 for b in blocks if b["kind"] == "table"),
            "image_ref_count": len(image_refs),
            "headings": collect_headings(blocks),
            "paragraphs": paragraphs,
            "truncated": truncated,
        }
        if tables_summary:
            result["tables_summary"] = tables_summary
        if image_refs:
            result["images"] = image_refs
        if next_offset is not None:
            result["next_offset"] = next_offset
        return result


def _natural_key(name: str) -> list[Any]:
    return [int(p) if p.isdigit() else p.lower() for p in re.split(r"(\d+)", name)]


def _caption_guess(blocks: list[dict[str, Any]], anchor_index: int) -> str | None:
    """Caption-styled or 'Figure N'-shaped paragraph next to the anchor."""
    for idx in (anchor_index + 1, anchor_index, anchor_index - 1):
        if idx < 0 or idx >= len(blocks):
            continue
        block = blocks[idx]
        if block["kind"] != "paragraph":
            continue
        text = block["text"].strip()
        if not text:
            continue
        if block["style"].lower() == "caption" or CAPTION_RE.match(text):
            return text
    return None


def _context_text(blocks: list[dict[str, Any]], anchor_index: int, limit: int = 240) -> str | None:
    """Nearest non-empty paragraph text around the anchor (anchor first)."""
    for idx in (anchor_index, anchor_index - 1, anchor_index + 1):
        if idx < 0 or idx >= len(blocks):
            continue
        block = blocks[idx]
        if block["kind"] != "paragraph":
            continue
        text = block["text"].strip()
        if text:
            return text[:limit]
    return None


def _prune_stale_extract_dirs(base: Path, max_age_seconds: float = 86400) -> None:
    import shutil
    import time

    cutoff = time.time() - max_age_seconds
    try:
        for child in base.iterdir():
            if child.is_dir() and child.stat().st_mtime < cutoff:
                shutil.rmtree(child, ignore_errors=True)
    except OSError:
        pass


def default_extract_output_dir() -> Path:
    base = Path(tempfile.gettempdir()) / "sylo-docx-extract"
    base.mkdir(parents=True, exist_ok=True)
    _prune_stale_extract_dirs(base)
    return Path(tempfile.mkdtemp(prefix="run-", dir=base))


def extract_docx_images(
    docx_path: Path,
    output_dir: Path | None = None,
) -> dict[str, Any]:
    """Extract all ``word/media/`` images, annotated with body anchors and captions.

    Every media image is extracted (matching the old template-docx-writer behavior);
    images referenced in the body additionally get ``block_index``,
    ``caption_guess``, ``context_text``, and ``alt_text`` so the agent knows
    which text talks about which picture. Unreferenced media (headers,
    footers, unused parts) are flagged.
    """
    docx_path = docx_path.resolve()
    if not docx_path.is_file():
        raise ValueError(f"DOCX not found: {docx_path}")

    out_dir = (
        Path(output_dir).resolve()
        if output_dir is not None
        else default_extract_output_dir()
    )
    out_dir.mkdir(parents=True, exist_ok=True)

    # Body walk first: media name → first anchor + caption/context.
    anchors: dict[str, dict[str, Any]] = {}
    with open_readable(docx_path) as (doc, _converted):
        blocks = walk_blocks(doc)
    for block in blocks:
        for ref in block["images"]:
            media = ref["media_name"]
            if media in anchors:
                anchors[media]["reference_count"] += 1
                continue
            info: dict[str, Any] = {
                "block_index": block["index"],
                "reference_count": 1,
            }
            if ref.get("alt_text"):
                info["alt_text"] = ref["alt_text"]
            caption = _caption_guess(blocks, block["index"])
            if caption:
                info["caption_guess"] = caption
            context = _context_text(blocks, block["index"])
            if context and context != caption:
                info["context_text"] = context
            anchors[media] = info

    saved: list[dict[str, Any]] = []
    with zipfile.ZipFile(docx_path, "r") as zf:
        media_files = sorted(
            (
                n
                for n in zf.namelist()
                if n.startswith(DOCX_MEDIA_PREFIX) and Path(n).suffix.lower() in IMAGE_EXTENSIONS
            ),
            key=lambda n: _natural_key(Path(n).name),
        )
        for entry in media_files:
            name = Path(entry).name
            ext = Path(name).suffix.lower()
            out_path = out_dir / name
            out_path.write_bytes(zf.read(entry))
            record: dict[str, Any] = {
                "media_name": name,
                "path": str(out_path),
                "format": ext.lstrip("."),
                "referenced_in_body": name in anchors,
            }
            if name in anchors:
                record.update(anchors[name])
            else:
                record["note"] = "Not referenced in document body (header/footer/theme or unused)."
            if ext in NON_VISION_FORMATS:
                record["format_note"] = f"{ext} may not be viewable by vision tools; convert to PNG if needed."
            saved.append(record)

    return {
        "docx_path": str(docx_path),
        "output_dir": str(out_dir),
        "ephemeral": output_dir is None,
        "image_count": len(saved),
        "referenced_count": sum(1 for s in saved if s["referenced_in_body"]),
        "images": saved,
    }
